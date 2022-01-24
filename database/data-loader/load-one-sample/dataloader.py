import os
import os
import json
import mysql.connector
import pandas as pd

from docx import Document
from mysql.connector import Error
from decimal import *
from numpy import FPE_UNDERFLOW, datetime64
from cmmidataloader import readCMMI
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl import load_workbook

class Config:
    rootSampleDirectory = ""
    hostName = ""
    userName  = ""
    userPassword= ""
    dbName = ""

class MMSampleSummaryData:
    sampleNo = 0
    sampleName = ""
    materialName = ""
    noSamples = 0
    description = ""
    RGBimage = ""

class MMSample:
    mmSampleId = 0
    sampleNo = 0
    date = 0
    waveLength = 0
    exposure = ""
    fNo = ""
    aoi = 0
    aoc= 0
    aos = 0
    mmImage = ""
    cmmiFilePath = ""

def parseSampleDescriptionFromWordDoc(sampleDirectoryPath, sampleNumber):
    # look for the word document in the folder and parse out the required information

    # build the file name
    wordDocFileName = f'DataSetInformation{sampleNumber}.docx'

    # build the path to the word document
    wordDocFilePath = os.path.join(sampleDirectoryPath, wordDocFileName)

    # load the word document
    doc = Document(wordDocFilePath)
    all_paras = doc.paragraphs
    print(len(all_paras))
    for para in all_paras:
        print(para.text)
        print("-------")
        
    sampleDescription = doc.paragraphs[17].text + doc.paragraphs[26].text
    sampleDescription = sampleDescription.strip()
    return sampleDescription

def createDBConnection(hostName, userName, userPassword, dbName):
    
    """Establish a connection to a mySQL server
        Parameters
        ----------
        host_name : str
        user_name : str
        user_password : str
        db_name : str
            
        Returns
        -------
        connection : connection object 
    """
    connection = None
    try:
        connection = mysql.connector.connect(
            host=hostName,
            user=userName,
            passwd=userPassword,
            database=dbName,
            auth_plugin = 'mysql_native_password'
        )
        print("MySQL Database connection successful")
    except Error as err:
        print(f"Error: '{err}'")

    return connection

def insertMMSample(dbConnection, mmSample):
    insertStmt = """INSERT INTO mm_samples (sample_no,date,wavelength,exposure,f_no,AOI,AOC,AOS,MMimage,cmmiFilePath) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
    data = (mmSample.sampleNo, mmSample.date, mmSample.wavelength, mmSample.exposure, mmSample.fNo, mmSample.aoi, mmSample.aoc, mmSample.aos,mmSample.MMimage,mmSample.cmmiFilePath)
    cursor = dbConnection.cursor()
    try:
        cursor.execute(insertStmt, data)
        dbConnection.commit()
        print("Query successful")
    except Error as err:
        print(f"Error: '{err}'")

# TODO: change to add the other samples, or create new excel sheet, or make excel sheet name a varibale and put into config
def parseExcelFile(config):
    # build the file name
    excelFileName = 'RGB950 2021 Sample List.xlsx'

    # build the excel file path
    excelFilePath = os.path.join(config.rootSampleDirectory, excelFileName)

    # TODO parse the summary data from the excel file
    wb = load_workbook(filename = excelFilePath)
    mmSampleSummaryDataRows = []
    
    sheet_ranges = wb['Final List']

    # load the sample no
    counter = 0
    for c in sheet_ranges['B']:
        print(c.value)
        if (counter >= 2): 
            mmSampleSummaryData = MMSampleSummaryData()
            mmSampleSummaryData.sampleNo = int(c.value)
            mmSampleSummaryDataRows.append(mmSampleSummaryData)
        counter = counter + 1

    # load the sample name
    counter = 0
    for c in sheet_ranges['C']:
        print(c.value)
        if (counter >= 2): 
            mmSampleSummaryData = mmSampleSummaryDataRows[counter - 2]
            mmSampleSummaryData.sampleName = c.value
        counter = counter + 1

    # load the number of samples
    counter = 0
    for c in sheet_ranges['D']:
        print(c.value)
        if (counter >= 2): 
            mmSampleSummaryData = mmSampleSummaryDataRows[counter - 2]
            mmSampleSummaryData.noSamples = c.value
        counter = counter + 1
    
    # load the material name
    counter = 0
    for c in sheet_ranges['F']:
        print(c.value)
        if (counter >= 2): 
            mmSampleSummaryData = mmSampleSummaryDataRows[counter - 2]
            mmSampleSummaryData.materialName = c.value
        counter = counter + 1
    return mmSampleSummaryDataRows


def insertMMSampleSummaryData(dbConnection, mmSampleSummaryData):
    insertStmt = """INSERT INTO material_samples (sample_no, sample_name, material_name, no_samples, description, RGBimage) VALUES (%s,%s,%s,%s,%s,%s)"""
    data = (mmSampleSummaryData.sampleNo,mmSampleSummaryData.sampleName,mmSampleSummaryData.materialName, mmSampleSummaryData.noSamples, mmSampleSummaryData.description, mmSampleSummaryData.RGBimage)
    cursor = dbConnection.cursor()
    try:
        cursor.execute(insertStmt,data)
        dbConnection.commit()
        print("Query successful")
    except Error as err:
        print(f"Error: '{err}'")

def loadCMMIFileData(mm, mmSample, dbConnection):
    image0 = mm[0]
    image1 = mm[1]
    image2 = mm[2]
    image3 = mm[3]
    image4 = mm[4]
    image5 = mm[5]
    image6 = mm[6]
    image7 = mm[7]
    image8 = mm[8]
    image9 = mm[9]
    image10 = mm[10]
    image11 = mm[11]
    image12 = mm[12]
    image13 = mm[13]
    image14 = mm[14]
    image15 = mm[15]

    for x in range(600):
        for y in range(600):
            m00 = image0[x][y]
            m01 = image1[x][y]
            m02 = image2[x][y]
            m03 = image3[x][y]
            m10 = image4[x][y]
            m11 = image5[x][y]
            m12 = image6[x][y]
            m13 = image7[x][y]
            m20 = image8[x][y]
            m21 = image9[x][y]
            m22 = image10[x][y]
            m23 = image11[x][y]
            m30 = image12[x][y]
            m31 = image13[x][y]
            m32 = image14[x][y]
            m33 = image15[x][y]

            insertStmt = """INSERT INTO cmmi_pixel_data (sample_no,pixel_x,pixel_y,m00,m10,m20,m30,m01,m11,m21,m31,m02,m12,m22,m32,m03,m13,m23,m33,AOI,AOC,wavelength) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
            data = (mmSample.sampleNo, x, y, m00, m01, m02, m03, m10, m11, m12, m13, m20, m21, m22, m23, m30, m31, m32, m33, mmSample.aoi, mmSample.aoc,mmSample.wavelength)
            cursor = dbConnection.cursor()
            try:
                cursor.execute(insertStmt,data)
                dbConnection.commit()
                print("Query successful")
            except Error as err:
                print(f"Error: '{err}'")

def loadCMMIFiles(cmmiDirectoryPath, dbConnection):
    # get all of the files in the cmmi directory
    for subdirs, dirs, files in os.walk(cmmiDirectoryPath):
        # loop over all of the files
        # TODO : pull out rgb image file path
        for fileName in files:
            # only process cmmi files
            if (fileName.__contains__(".cmmi")):

                # parse the folder name for the sample data
                mmSample = MMSample()
                mmSample.date = fileName.rsplit("_", 7)[0]
                mmSample.sampleNo = fileName.rsplit("_", 7)[1]
                mmSample.wavelength = fileName.rsplit("_", 7)[2]
                mmSample.exposure = fileName.rsplit("_", 7)[3]
                mmSample.fNo = fileName.rsplit("_", 7)[4]
                mmSample.aoi = int(fileName.rsplit("_", 7)[5])
                mmSample.aoc = fileName.rsplit("_", 7)[6]
                mmSample.aoc = int(mmSample.aoc.rsplit(".",2)[0])
                mmSample.aos = mmSample.aoc - mmSample.aoi
                mmSample.MMimage = cmmiDirectoryPath
                mmSample.cmmiFilePath = cmmiDirectoryPath

                # insert the sample into the database
                insertMMSample(dbConnection, mmSample)

                # build the path to the cmmi file that we are processing
                cmmiFilePath = os.path.join(cmmiDirectoryPath, fileName)

                # parse the cmmi file
                mm = readCMMI(cmmiFilePath)

                # load the cmmi file
                # table no longer in use
                # loadCMMIFileData(mm, mmSample, dbConnection)

def loadSample(sampleDirectoryPath, config):
    # determine the sample number from the path
    sampleNo = os.path.basename(sampleDirectoryPath)
    print (sampleNo)

    # create a db connection
    dbConnection = createDBConnection(config.hostName, config.userName, config.userPassword, config.dbName)

    # parse the excel file
    mmSampleSummaryDataRows = parseExcelFile(config)

    # find sample data in data rows
    for mmSampleSummaryData in mmSampleSummaryDataRows:
        if (mmSampleSummaryData.sampleNo == sampleNo):
            break

    # determine the sample description
    mmSampleSummaryData.description = parseSampleDescriptionFromWordDoc(sampleDirectoryPath, sampleNo)

    # insert sample summary data
    insertMMSampleSummaryData(dbConnection, mmSampleSummaryData)

    # loop over wave length directory in the sample directory
    for subdirs, dirs, files in os.walk(sampleDirectoryPath):
        for waveLengthDirectoryName in dirs:

            # build the path to the wavelength directory 
            cmmiDirectoryPath = os.path.join(sampleDirectoryPath, waveLengthDirectoryName, "cmmi")

            # load the wave length
            loadCMMIFiles(cmmiDirectoryPath, dbConnection)



    